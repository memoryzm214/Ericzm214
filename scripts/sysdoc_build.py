# -*- coding: utf-8 -*-
import docx, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.6)
sec.top_margin = sec.bottom_margin = Cm(2.4)

st = doc.styles['Normal']
st.font.name = 'Times New Roman'; st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def setfont(run, cn='宋体', en='Times New Roman', size=10.5, bold=False, color=None):
    run.font.name = en; run.font.size = Pt(size); run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    if color: run.font.color.rgb = color

def P(text='', size=10.5, cn='宋体', bold=False, align='both', first=True,
      before=0, after=6, line=18, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(before), Pt(after)
    pf.line_spacing = Pt(line)
    pf.alignment = {'both': WD_ALIGN_PARAGRAPH.JUSTIFY, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'left': WD_ALIGN_PARAGRAPH.LEFT}[align]
    if first: pf.first_line_indent = Pt(size * 2)
    if text:
        setfont(p.add_run(text), cn=cn, size=size, bold=bold, color=color)
    return p

def H(text, size=13, before=14, after=8):
    P(text, size=size, cn='黑体', bold=False, align='left', first=False,
      before=before, after=after, line=size*1.6)

def H2(text, size=11, before=10, after=5):
    P(text, size=size, cn='黑体', bold=False, align='left', first=False,
      before=before, after=after, line=size*1.6)

def rule():
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    b = pPr.makeelement(qn('w:pBdr'), {}); pPr.insert(0, b)   # pBdr 须排在 spacing 之前
    bot = b.makeelement(qn('w:bottom'), {qn('w:val'):'single', qn('w:sz'):'6',
                                          qn('w:space'):'1', qn('w:color'):'888888'})
    b.append(bot)

# ================= 封面标题 =================
P('2026年度湖北省社会科学基金一般项目', size=16, cn='黑体', align='center',
  first=False, before=18, after=4, line=26)
P('网上申报系统填报内容', size=16, cn='黑体', align='center', first=False, after=14, line=26)
P('（项目选题 · 内容简介 · 主要参考文献 · 经费预算）', size=10.5, align='center',
  first=False, after=4, line=18, color=RGBColor(0x55,0x55,0x55))
rule()

# ================= 填写说明 =================
H('填写说明', size=12, before=4)
P('一、本文件依据《高校共享道路慢行主体的安全认知偏差与行为引导研究》研究报告第二轮校订定稿版（103,103字符，图11幅、表41张、参考文献87条）拟写，各项文字可直接复制粘贴至申报系统对应文本框。')
P('二、系统路径为“一般项目管理→项目申报”，且须填报经费预算，故判断本次为立项申报而非成果奖申报。因此“项目内容简介”按项目申报的规范文体撰写，即以拟开展的研究为主线，已完成的工作作为前期基础呈现。若贵方实际情形为成果类申报，需将时态改为已完成式，请告知。')
P('三、经费方案为按课题实际研究活动测算的建议值，合计 20000.00 元。湖北省社会科学基金对各科目的比例限制（尤其是间接费用、劳务费与专家咨询费的上限）以当年度申报公告及经费管理办法为准，提交前请对照核定。')
P('四、字数按系统提示控制：项目选题约 1500 字，内容简介约 2500 字，参考文献 20 项。文中数据均取自研究报告，与报告正文严格一致。第五节另附可供申报材料使用的图表清单，系统文本框不支持插图，该节供制作申报书附件时取用。')

# ================= 一、基本信息 =================
H('一、项目名称与关键词')
P('项目名称：高校共享道路慢行主体的安全认知偏差与行为引导研究', first=False)
P('关 键 词：高校共享道路；慢行主体；安全认知偏差；交通冲突；行为决策模型；安全引导', first=False)

# ================= 二、项目选题 =================
H('二、（3）项目选题')
P('（要点提示：1500字左右，课题国内外研究现状述评；选题的意义）', size=9.5, first=False,
  after=8, color=RGBColor(0x88,0x88,0x88))

H2('（一）课题国内外研究现状述评')
P('共享空间理念主张通过弱化道路边界、减少显性管制，促使不同交通参与者以协商方式分配路权。已有研究对该理念的演进作了系统梳理，指出它并非单一设计手法，而是由“完全共享”到“部分分隔”的连续谱系。但该理念在高密度场景中的适用性持续存在争议：有研究指出，缺乏明确路权提示时，视障者等弱势使用者会因无法获取必要的定向信息而回避使用该空间；风险补偿理论与风险稳态理论则提示，安全条件的改善有可能被更为激进的行为所抵消。上述争议共同指向一个前提性问题，即共享空间的有效性依赖于使用者能否准确感知环境的真实风险。')
P('在冲突识别方面，交通冲突技术以碰撞时间、后侵入时间等替代安全指标，实现了事故发生之前的安全评价；近年研究结合视频轨迹提取与机器学习方法，已能较为精确地刻画冲突的时空分布与行为类型。但已有研究亦指出，上述指标体系多源自机动车交通情境，而行人与骑行者轨迹自由度高、速度低、方向变化频繁，其规避行为往往在指标达到临界值之前即已完成，直接移用会系统性地低估慢行冲突的严重程度。')
P('在主观安全感知方面，自行车服务水平、交通压力等级、步行适宜性指数等方法，将“是否感到安全”这一主观判断锚定到可测量的物理属性之上。面向校园的研究进一步发现，客观审计结果与使用者主观评分在部分路段出现明显分歧，客观条件较好的路段未必获得较高评价。这一发现具有重要提示意义，但既有研究多将其作为方法学上的信度问题加以处理，鲜有将这种偏离本身确立为研究对象。')
P('在行为机理方面，计划行为理论、合理行为理论与技术接受模型被广泛用于解释交通违规意图的形成。有研究针对校园电动自行车使用者开展结构方程建模，发现感知有用性与感知易用性对违规行为的影响最为显著，提示提升合规行为的便利性较之加大惩罚力度更为有效。风险认知研究则表明，个体对风险的判断受可控性、熟悉度等定性特征的系统影响，并普遍依赖可得性启发式，因而所产生的偏差并非随机误差，而是具有稳定结构、可被识别与预测的现象。')
P('在引导策略方面，国内外已提出工程、教育、管理与技术四类干预路径。但这些措施与前述行为机理研究之间往往缺乏明确的推导关系，多为经验判断或案例移植的结果；措施为何有效、对哪一类使用者有效、经由何种心理路径起作用，通常未被交代，这既限制了措施的可迁移性，也使其效果难以预期。')
P('综合而言，既有研究之间存在三处尚未打通的环节：其一，冲突的客观识别与参与者心理机制的解释各自成立，却尚未形成相互印证的证据链；其二，主观感知与客观风险之间的系统性偏差未被作为独立对象加以测度；其三，由理论发现通向设计实践的转化路径不明。本课题正是针对上述三处环节而设。')

H2('（二）选题的意义')
P('理论意义体现在三个方面。第一，将安全认知偏差由方法学上的信度问题转化为具有实质意义的研究对象，并为其建立可计算的量化指标，使“主观感知与客观风险之间相差多少、在何处相差”成为可测量、可比较的问题。第二，以行人与骑行者所构成的慢行主体作为统一分析单元，揭示同一空间中的两类主体并非遵循同一套心理逻辑，从而为差异化引导提供理论依据，也在一定程度上解释了面向全体使用者的统一化管理措施何以收效有限。第三，将行为科学的解释性框架与设计学的生成性方法相衔接，形成“识别—解释—生成—验证”的完整闭环，为设计干预类研究提供一种可复用的研究范式。')
P('实践意义在于回应校园交通治理的特殊约束。高校管理部门并不具有道路交通执法权，可用手段主要是宣传教育、志愿劝导与设施改造，依靠处罚形成威慑的路径在校园中难以奏效，而通过环境设计引导行为的路径则具有更大的施展空间。本课题所形成的改造优先级依据、可跨校迁移的安全引导系统方案，以及供高校自行开展的安全自评估流程，可直接服务于湖北省高校的平安校园与绿色校园建设。湖北省高等教育资源富集，校园共享道路的安全问题具有相当覆盖面，相应成果亦具备在省域推广的现实基础，对提升师生日常安全感与校园环境品质具有现实价值。')

doc.add_page_break()

# ================= 三、内容简介 =================
H('三、（4）项目内容简介', before=0)
P('（要点提示：2500字左右，课题研究的基本思路和方法，主要观点和创新点等）', size=9.5,
  first=False, after=8, color=RGBColor(0x88,0x88,0x88))

H2('（一）研究的基本思路')
P('本课题聚焦高校共享道路中慢行主体的交通安全问题。所称慢行主体，特指共享道路上的行人与骑行者两类交通参与者，其中骑行者涵盖使用自行车与电动自行车者；所称共享道路，指未在不同交通方式之间作物理分隔、由上述两类主体共同使用的校园道路。之所以以慢行主体而非单一交通方式立题，是因为课题所关注的安全认知偏差与交通冲突，均产生于两类主体的相互作用之中：行人的风险感知取决于骑行者的速度与轨迹，骑行者的通行判断亦受制于行人的分布与避让预期，脱离任何一方都无法对冲突的形成作出完整解释。')
P('课题的基本判断是：此类空间的交通冲突，既来自路权的物理竞争，也来自慢行主体在安全认知上的系统性偏差——同一路段的风险，在行人与骑行者眼中往往并不相同，而两者的判断又都可能偏离客观实际。据此，课题以安全认知偏差为枢纽，遵循“特征识别—模型构建—需求识别—策略生成”的研究路径，构建四项循序渐进的子任务：慢行主体交通冲突的特征识别；安全认知与行为意图的作用机理研究；主客观认知偏差的测度与需求分析；安全引导策略的生成与验证。')
P('四项子任务之间是逐级供给关系而非彼此独立的并列罗列。子任务①以视频观测界定冲突的空间边界与行为边界，其结果为子任务②确定情境背景；子任务②揭示行为意图的形成机理，所识别的关键变量为子任务③的场景变量设定提供依据；子任务③以主客观对照量化认知偏差，并经访谈与工作坊将偏差转化为可操作的设计需求；子任务④以需求清单为输入生成设计方案，并回到量化评估完成验证。')

H2('（二）研究方法')
P('课题采用混合方法研究范式，以探索性时序设计组织定量与定性方法。之所以作此选择，是由研究问题本身的性质决定的：冲突在何时、何地、以何种形式发生，需要客观观测与统计描述才能回答；参与者为何如此行动、其判断依据为何，则需要理解意义与还原过程才能回答，二者缺一不可。')
P('子任务①采用视频观测与交通冲突技术。不以事故统计为基础，是因为校园内经正式登记的事故极少且大量轻微冲突不被记录，难以获得足以支撑分析的样本量；而冲突技术的前提正是大量轻微交互与少量事故之间存在可推断的统计关联。观测点位经实地踏勘与专家咨询共同确定；各点位摄像设备分别建立单应性矩阵以消除俯角拍摄的透视畸变，轨迹经平滑后换算速度并计算碰撞时间；人工复核环节安排双人独立标注并计算一致性系数。考虑到校园慢行交通速度低、轨迹变化大、注意力分散频繁，课题对碰撞时间的分级阈值作针对性下调，以免大量实际已构成危险的交互被归入低等级。')
P('子任务②采用结构方程模型。因安全态度、主观规范、知觉行为控制等均属不可直接观测的潜在构念且构念之间存在多重中介关系，需同时估计测量模型与结构模型，较之一般回归分析更为适宜。量表在整合计划行为理论、合理行为理论与技术接受模型的基础上开发，经专家内容效度评定、预试探索性因子分析与正式验证性因子分析后使用；在开展行人与骑行者的多群组比较之前，先行完成组态、度量与标量三个层级的测量不变性检验，以确保组间路径差异可被解释为心理机制的实质差异，而非测量层面的人为产物。')
P('子任务③采用人工智能生成的标准化场景作为实验刺激材料。不使用实地照片，是因为实地照片各环境变量高度耦合，难以判断评分差异由哪一要素引起；而生成式方法可在保持其他条件不变的前提下逐一改变单个变量，从而使主效应的估计成为可能。课题以正交设计从完整因子空间中筛选代表性场景，构建客观安全指数与认知偏差系数的量化对照体系，并结合深度访谈与参与式设计工作坊提炼设计需求。')
P('子任务④采用参与式设计与案例研究相结合的路径生成方案，以德尔菲法与模糊综合评价择优；验证环节采用主观评价与客观测算并行的策略，以系统可用性量表和出声思维法考察可理解性与易用性，以自行车服务水平模型和路径使用者舒适度评估测算通行品质的改善幅度。')

H2('（三）主要观点')
P('第一，高校共享道路的交通冲突是认知、行为与环境三者耦合作用的结果，而安全认知偏差正是三者之间的连接点。环境要素通过影响信息的可获取性塑造使用者的风险判断，风险判断继而经由各自的心理路径转化为行为意图，行为落实之后又反过来改变环境中的交通状态，形成新的冲突情境。')
P('第二，行人与骑行者遵循不同的行为决策机理。前期研究显示，行人的决策以风险感知驱动安全态度、再由态度决定行为意图；骑行者则以设施的易用性支撑控制感、再由控制感决定行为意图。这一差异意味着安全引导不宜采取无差别的统一形式，而应针对两类主体各自的心理路径分别设计。')
P('第三，认知偏差在空间上呈现规律性分布。前期研究显示，视距受限的弯道路段风险被系统性低估，而建筑出入口的风险则被高估；尤为值得注意的是，视距条件在客观上是事故预防的关键因素，却未能进入主观安全感知的显著预测模型，这一“客观关键、主观忽略”的错位，正是设计干预应当瞄准的靶点。')
P('第四，在校园这一缺乏执法权的柔性治理场域中，提升合规行为的便利性较之加大违规成本更为有效；单纯增加标识数量对冲突率的改善有限，其原因在于新增信息未能落在使用者实际形成判断的决策点上，反而增加视觉负荷。')

H2('（四）创新点')
P('理论创新：将安全认知偏差由方法学上的信度问题确立为可测度的研究对象，建立客观安全指数与认知偏差系数的量化对照体系，使主客观偏离成为可比较、可定位的变量，突破了既有研究将主客观不一致视为测量误差的处理方式。')
P('方法创新：以慢行主体为统一分析单元，通过多群组结构方程分析分别刻画行人与骑行者的决策机理，并以测量不变性检验保证组间比较的可靠性；同时引入生成式场景实验，在保持其他条件不变的前提下逐一操控环境变量，解决了实地照片中变量高度耦合、主效应无法估计的难题。')
P('实践创新：构建由认知偏差直接推导设计需求、再由需求生成可实施引导方案的转化路径，并形成包含色彩、形态与触觉三重冗余编码的引导系统设计规范。经测算，方案所用冷蓝与暖橙两色的明度对比度低于图形要素通常要求的水平，单凭色彩不足以支持色觉障碍使用者的区分，故须辅以白色边界实线与地面纹理差异——这一由计算得出的约束，构成了冗余编码原则的直接依据。')

doc.add_page_break()

# ================= 四、参考文献 =================
H('四、（5）主要参考文献（20项）', before=0)
REFS = [
'李琳, 叶宇, 陈泳. 城市步行安全及其环境影响要素研究综述与展望[J]. 风景园林, 2025, 32(2): 86-94.',
'吕能超, 王玉刚, 周颖, 等. 道路交通安全分析与评价方法综述[J]. 中国公路学报, 2023, 36(4): 183-201.',
'刘孟歆, 秦华, 岳晨, 等. 右转车辆与过街行人交互过程的影响因素研究[J]. 包装工程, 2023, 44(12): 118-125.',
'裴玉龙, 龙钰, 马丹. 交通安全意识对非机动车骑行者危险骑行行为的影响研究[J]. 交通信息与安全, 2024, 42(1): 49-58+66.',
'张乐, 汤晓敏. 可供性视角下大学老校区公共空间评价与更新设计指引——以上海交通大学徐汇校区为例[J]. 中国园林, 2025, 41(5): 115-122.',
'陈雅楠, 赵晓华, 李佳, 等. 基于科学知识图谱的道路交叉口安全设施设计综述及范式研究[J]. 北京工业大学学报, 2024, 50(12): 1501-1520.',
'ZHANG Z, FISHER T, WANG H. Walk Score, environmental quality and walking in a campus setting[J]. Land, 2023, 12(4): 732.',
'CHEN H, GUO Y, LI L. Promoting sustainable mobility on campus: uncovering the behavioral mechanisms behind non-compliant e-bike use among university students[J]. Sustainability, 2025, 17(15): 7147.',
'MIAO L, LIU F, DENG Y. Analysis of traffic conflicts on slow-moving shared paths in Shenzhen, China[J]. Sustainability, 2025, 17(9): 4095.',
'KELLSTEDT K, SPENGLER J O, MADDOCK J E. Comparing perceived and objective measures of bikeability on a university campus: a case study[J]. SAGE Open, 2021, 11(2): 21582440211018685.',
'VON STÜLPNAGEL R, RINTELEN H. A matter of space and perspective: cyclists\', car drivers\' and pedestrians\' assumptions about subjective safety in shared traffic situations[J]. Transportation Research Part A: Policy and Practice, 2024, 179: 103941.',
'HAMILTON-BAILLIE B. Shared space: reconciling people, places and traffic[J]. Built Environment, 2008, 34(2): 161-181.',
'KARNDACHARUK A, WILSON D J, DUNN R. A review of the evolution of shared (street) space concepts in urban environments[J]. Transport Reviews, 2014, 34(2): 190-220.',
'JOHNSSON C, LAURESHYN A, DE CEUNYNCK T. In search of surrogate safety indicators for vulnerable road users: a review of surrogate safety indicators[J]. Transport Reviews, 2018, 38(6): 765-785.',
'HAYWARD J C. Near-miss determination through use of a scale of danger[J]. Highway Research Record, 1972, 384: 24-34.',
'AJZEN I. The theory of planned behavior[J]. Organizational Behavior and Human Decision Processes, 1991, 50(2): 179-211.',
'DAVIS F D. Perceived usefulness, perceived ease of use, and user acceptance of information technology[J]. MIS Quarterly, 1989, 13(3): 319-340.',
'TVERSKY A, KAHNEMAN D. Judgment under uncertainty: heuristics and biases[J]. Science, 1974, 185(4157): 1124-1131.',
'SLOVIC P. Perception of risk[J]. Science, 1987, 236(4799): 280-285.',
'THALER R H, SUNSTEIN C R. Nudge: improving decisions about health, wealth, and happiness[M]. New Haven: Yale University Press, 2008.',
]
for i, r in enumerate(REFS, 1):
    p = P(f'[{i}] {r}', first=False, after=3, line=16)
    p.paragraph_format.left_indent = Pt(22)
    p.paragraph_format.first_line_indent = Pt(-22)


doc.add_page_break()

# ================= 五、可供申报使用的图表 =================
H('五、可供申报使用的图表', before=0)
P('系统的四个文本框仅接受纯文字，下列图表供制作申报书正文与附件时取用，均出自研究报告，编号与报告一致。', size=9.5, first=False, after=8, color=RGBColor(0x88,0x88,0x88))

H2('（一）建议优先选用的图')
FIGS = [
 ('图1', '本课题技术路线图', '置于“研究思路与方法”，一图说明四项子任务的逐级供给关系'),
 ('图2', '校园共享道路12个观测点', '置于“研究基础”，说明实证对象与观测点覆盖面'),
 ('图5、图6', '行人／骑行者行为决策模型', '置于“主要观点”，直观呈现两类主体的路径差异'),
 ('图9', '安全引导系统设计原型（地面标识系统）', '置于“预期成果”，展示设计落点'),
 ('图10', '交叉口部署预警标志设计原型', '与图9配合，说明冲突点的针对性处理'),
 ('图11', '立面动态提示装置设计原型', '说明“静态与动态”互补的系统构成'),
]
t1 = doc.add_table(rows=1, cols=3)
t1.style = 'Table Grid'; t1.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['图号', '图名', '建议用途']):
    c = t1.rows[0].cells[j]; c.width = [Cm(2.0), Cm(6.4), Cm(7.4)][j]; c.text = ''
    r = c.paragraphs[0].add_run(h); setfont(r, cn='黑体', size=10)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for a, b, cc in FIGS:
    row = t1.add_row()
    for j, v in enumerate([a, b, cc]):
        c = row.cells[j]; c.width = [Cm(2.0), Cm(6.4), Cm(7.4)][j]; c.text = ''
        pp = c.paragraphs[0]; pp.paragraph_format.space_before = Pt(2); pp.paragraph_format.space_after = Pt(2)
        setfont(pp.add_run(v), size=9.5)
        if j == 0: pp.alignment = WD_ALIGN_PARAGRAPH.CENTER

H2('（二）建议优先选用的表')
TBLS = [
 ('表7', '校园共享道路观测点', '说明观测点位的空间类型覆盖'),
 ('表11、表12', '冲突事件空间分布统计、冲突行为模式统计', '支撑“冲突集聚”与“人为因素主导”两项判断'),
 ('表13', '交通环境与冲突行为的相关性统计', '给出道路宽度、标识完整度、视距条件的相关系数'),
 ('表23', '多群组分析路径系数比较', '支撑“两类主体机理不同”这一核心观点'),
 ('表30、表31', '不同场景类型的认知偏差统计、客观因素回归分析', '支撑“弯道低估、出入口高估、视距被忽略”'),
 ('表38', '安全引导系统设计参数规范', '体现成果的可施工性与可复现性'),
 ('表41', '设计方案实施前后的量化评估对比', '给出BLOS、PUCE等指标的改善幅度'),
]
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'; t2.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(['表号', '表名', '建议用途']):
    c = t2.rows[0].cells[j]; c.width = [Cm(2.0), Cm(6.4), Cm(7.4)][j]; c.text = ''
    r = c.paragraphs[0].add_run(h); setfont(r, cn='黑体', size=10)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for a, b, cc in TBLS:
    row = t2.add_row()
    for j, v in enumerate([a, b, cc]):
        c = row.cells[j]; c.width = [Cm(2.0), Cm(6.4), Cm(7.4)][j]; c.text = ''
        pp = c.paragraphs[0]; pp.paragraph_format.space_before = Pt(2); pp.paragraph_format.space_after = Pt(2)
        setfont(pp.add_run(v), size=9.5)
        if j == 0: pp.alignment = WD_ALIGN_PARAGRAPH.CENTER

H2('（三）图例：观测点分布与设计原型')
import os as _os
def figure(pathname, caption, width_cm=15.2):
    if not _os.path.exists(pathname):
        return
    pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(6); pp.paragraph_format.space_after = Pt(2)
    pp.add_run().add_picture(pathname, width=Cm(width_cm))
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(10)
    setfont(cp.add_run(caption), size=9.5, color=RGBColor(0x55,0x55,0x55))

figure('/home/user/Ericzm214/配图备选/图2_校园观测点_淡绿色版.png', '图2  校园共享道路12个观测点（淡绿色系重绘版）')
figure('/home/user/Ericzm214/配图备选/图9/图9-C.png', '图9  安全引导系统设计原型（要素标注版，备选方案9-C）')
figure('/home/user/Ericzm214/配图备选/图11/图11-D.png', '图11  立面动态提示装置屏幕界面（主信息区70%／侧栏30%，备选方案11-D）')
P('上列三幅为本次重绘的备选版本，另有图9、图10、图11 各6张备选存于《配图备选》目录，可按版面需要替换。', size=9.5, first=False, color=RGBColor(0x88,0x88,0x88))


doc.add_page_break()

# ================= 六、经费 =================
H('六、6. 经费信息', before=0)
P('单位：元（填写数字，保留两位小数）。合计 20000.00 元。', size=9.5, first=False,
  after=8, color=RGBColor(0x88,0x88,0x88))

BUDGET = [
 ('资料费', '2000.00', '中外文数据库检索与文献传递、专业图书与标准文本购置'),
 ('会议费/差旅费/国际合作与交流费', '2500.00', '省内高校实地调研差旅、学术会议交流与阶段成果研讨'),
 ('设备费', '1500.00', '视频采集设备及存储介质、激光测距仪与照度计等测量器材租用及耗材'),
 ('专家咨询费', '3000.00', '德尔菲法两轮专家咨询、客观指标赋权、量表评定与设计方案评审'),
 ('劳务费', '5000.00', '调查员与视频标注员劳务、访谈与工作坊参与者补助、研究生助理'),
 ('印刷出版费', '3000.00', '论文发表版面费、研究报告与设计图册印制'),
 ('其它支出', '1000.00', '问卷平台与生成式图像平台服务、设计原型打样与材料'),
 ('间接费用', '2000.00', '依托单位管理费用及绩效支出，占直接费用的11.1%'),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ['经费科目', '金额（元）', '测算依据']
widths = [Cm(5.4), Cm(2.6), Cm(7.8)]
for j, h in enumerate(hdr):
    c = tbl.rows[0].cells[j]; c.width = widths[j]
    c.text = ''
    r = c.paragraphs[0].add_run(h); setfont(r, cn='黑体', size=10, bold=False)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for name, amt, basis in BUDGET:
    row = tbl.add_row()
    for j, v in enumerate([name, amt, basis]):
        c = row.cells[j]; c.width = widths[j]; c.text = ''
        r = c.paragraphs[0].add_run(v); setfont(r, size=9.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
        c.paragraphs[0].paragraph_format.space_after = Pt(2)
row = tbl.add_row()
for j, v in enumerate(['合计', '20000.00', '—']):
    c = row.cells[j]; c.width = widths[j]; c.text = ''
    r = c.paragraphs[0].add_run(v); setfont(r, cn='黑体', size=10, bold=False)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

P('', after=4, first=False)
P('说明：上表“测算依据”一列供填报时参考，系统仅需填入金额。各科目比例限制以当年度申报公告及湖北省社会科学基金经费管理办法为准，提交前请对照核定；若公告对间接费用或劳务费设有上限而本方案超出，可在保持合计不变的前提下于科目之间调整。', size=9.5)

# ================= 六、待确认 =================
H('七、提交前需确认的事项')
for t in [
 '1. 经费各科目的比例限制，以当年度申报公告与湖北省社会科学基金经费管理办法为准核定；若公告对间接费用或劳务费设有上限而本方案超出，可在保持合计 20000.00 元不变的前提下于科目之间调整。',
 '2. 研究报告封面首行的课题批次标注与日期仍为空，需按本次申报材料的用途填写；封面“课题负责人”“申报单位”两栏亦待填。',
 '3. 报告中涉及湖北省高校规模的表述目前为定性描述，如在申报材料中引用，建议补入《湖北省教育事业发展统计公报》的具体数据与年份。',
 '4. 报告第1章提及的湖北省教育强省与平安校园建设相关文件，需补入准确名称与发文年份。',
 '5. 本文件参考文献 20 项自研究报告 87 条中选取，含中文文献 6 条，覆盖共享空间理论、冲突识别方法、主观安全评价、行为机理与行为引导五个方面；提交前请对照原文复核卷、期、页码。省社科基金评审通常期待更充分的中文文献覆盖，建议在知网另行补充后同步更新本清单。',
 '6. 你在最新一版报告中删除了正文中“图4  校园共享道路5个热点区域的具体交通状况”的图注，但图目录中该条仍在，正文4.1.2节亦写有“如图3、图4所示”。若属误删请补回，若确需删除则需同步修改图目录与正文表述。',
 '7. 附录F已按你的修改去掉“职称层次”一列，正文对应的专家构成表述已随之调整为仅列专业领域。',
 '8. 图9、图10、图11 与图2 的重绘版本各有备选，尚未替换进研究报告正文；确定选用哪一版后可一并替换并更新图目录。',
]:
    P(t, first=False, after=4)

_z = doc.settings.element.find(qn('w:zoom'))
if _z is not None and _z.get(qn('w:percent')) is None:
    _z.set(qn('w:percent'), '100')

doc.save('/home/user/Ericzm214/申报系统填报内容.docx')

# 字数核验
import docx as _d
dd = _d.Document('/home/user/Ericzm214/申报系统填报内容.docx')
txt = [p.text for p in dd.paragraphs]
def seg(a, b):
    ia = next(i for i,t in enumerate(txt) if t.startswith(a))
    ib = next(i for i,t in enumerate(txt) if t.startswith(b))
    return sum(len(re.sub(r'\s','',t)) for t in txt[ia+1:ib])
print('项目选题字数：', seg('二、（3）项目选题', '三、（4）项目内容简介'))
print('内容简介字数：', seg('三、（4）项目内容简介', '四、（5）主要参考文献'))
print('参考文献条数：', sum(1 for t in txt if re.match(r'^\[\d+\] ', t)))
